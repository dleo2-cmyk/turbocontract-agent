# -*- coding: utf-8 -*-
"""
TurboContract Document Analysis Agent — Prototype
==================================================
Reads .docx documents, detects variables/blocks via cross-document comparison
and LLM (few-shot via OpenRouter), then writes results to Excel "Анализ сложности" sheet.

Requirements:
    pip install openai python-docx openpyxl

Usage:
    python agent_prototype.py
    python agent_prototype.py --docs "path/to/docs" --excel "path/to/table.xlsx"
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import io
from collections import defaultdict
from pathlib import Path

# Загружаем .env если есть (локальная разработка)
_env_file = Path(__file__).parent / ".env"
if _env_file.exists():
    for _line in _env_file.read_text(encoding="utf-8").splitlines():
        _line = _line.strip()
        if _line and not _line.startswith("#") and "=" in _line:
            _k, _v = _line.split("=", 1)
            os.environ.setdefault(_k.strip(), _v.strip())
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Optional, List, Tuple, Dict, Any

from openai import OpenAI
from docx import Document
from docx.table import Table
import openpyxl

# ─────────────────────────────────────────────
# Fix Windows console encoding
# ─────────────────────────────────────────────
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

# ─────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────
DEFAULT_DOCS_DIR = r"C:\Users\dleo2\Downloads\Договор"
DEFAULT_EXCEL = r"C:\Users\dleo2\Downloads\Таблица_для_расчета_сложности_шаблонов.xlsx"

# API-ключ берётся из переменной среды или .env файла (не хранить в коде!)
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "")

OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"

# Модель на OpenRouter — можно менять
# Варианты: "anthropic/claude-opus-4-5", "anthropic/claude-sonnet-4-5",
#            "google/gemini-2.0-flash-001", "openai/gpt-4o"
MODEL = "openai/gpt-4o"

# ─────────────────────────────────────────────
# Regex patterns for deterministic variable counting
# ─────────────────────────────────────────────
# Matches [МАРКЕР:text] tags injected by _annotate_para()
_MARKER_RE = re.compile(r'\[МАРКЕР:([^\]]+)\]')
# Matches underscore placeholders: ___ (3 or more)
_UNDERSCORE_RE = re.compile(r'_{3,}')
# Минимальный стеммер для русского — убирает падежные окончания
# чтобы «заказчика» и «заказчику» давали один контекстный ключ
_RU_SUFFIXES = (
    'ями', 'ами', 'ого', 'его', 'ому', 'ему', 'ой', 'ей',
    'ом', 'ем', 'ах', 'ях', 'ые', 'ие', 'ую', 'юю',
    'ью', 'ьи', 'ья', 'ов',
    'а', 'я', 'е', 'у', 'ю', 'ы', 'и', 'ь', 'й',
)

def _stem(word: str) -> str:
    for suf in _RU_SUFFIXES:
        if word.endswith(suf) and len(word) - len(suf) >= 3:
            return word[:-len(suf)]
    return word

# Service tables (реквизиты/подписи) — не считать как содержательные таблицы.
# ТОЛЬКО слова уникальные для таблиц реквизитов — НЕ ИНН/КПП/БИК,
# они могут встречаться и в содержательных таблицах (спецификации с реквизитами подрядчика).
# М.П. / место печати — практически исключительно в блоках подписей.
# Банковские р/с, к/с — маркеры таблицы реквизитов сторон.
_SERVICE_TABLE_RE = re.compile(
    r'реквизит'
    r'|м\.п\.|место\s+печати'
    r'|р/с|к/с'                        # банковские счета — только в таблицах реквизитов
    r'|расчётный\s+счёт|корреспондентский\s+счёт'
    # Любые две стороны договора в ОДНОЙ строке = заголовок таблицы подписей
    r'|(?:заказчик|продавец|арендодатель|лицензиар|займодавец|залогодатель|'
     r'поставщик|исполнитель|покупатель|подрядчик|арендатор|лицензиат|заёмщик|залогодержатель)'
     r'[^|\n]{0,60}\|[^|\n]{0,60}'
     r'(?:заказчик|продавец|арендодатель|лицензиар|займодавец|залогодатель|'
     r'поставщик|исполнитель|покупатель|подрядчик|арендатор|лицензиат|заёмщик|залогодержатель)'
    # «Арендодатель» в кавычках-ёлочках как отдельная ячейка = заголовок блока подписей
    # (даже если Арендатор/Подрядчик стоит в другой строке/ячейке той же таблицы)
    r'|«(?:арендодатель|арендатор|заказчик|подрядчик|исполнитель|'
     r'покупатель|поставщик|продавец|лицензиар|лицензиат|займодавец|'
     r'залогодатель|залогодержатель)»'
    # Таблица-подпись одной стороны: первая строка = «ИСПОЛНИТЕЛЬ:» (без пары в той же строке)
    r'|^(?:исполнитель|заказчик|подрядчик|арендодатель|арендатор|'
     r'поставщик|покупатель|продавец|лицензиар|лицензиат|'
     r'займодавец|залогодержатель)\s*:',
    re.IGNORECASE
)

# Appendix section headings embedded inside the main document body.
# Text like «ПРИЛОЖЕНИЕ №1» / «Приложение 2 к договору» signals that all
# following paragraphs and tables belong to an in-document appendix and
# must NOT be counted as main contract content.
_APPENDIX_HEADING_RE = re.compile(
    r'^\s*приложени[еия]\s*(?:№\s*)?\s*\d',
    re.IGNORECASE
)

# ─────────────────────────────────────────────
# Data structures
# ─────────────────────────────────────────────
@dataclass
class DocInfo:
    """Raw content extracted from one .docx file."""
    path: str
    name: str
    paragraphs: list[str] = field(default_factory=list)
    appendix_paragraphs: list[str] = field(default_factory=list)  # текст параграфов в приложениях
    tables_count: int = 0
    content_tables_count: int = 0           # tables excluding реквизиты/подписи
    content_tables_content: list[str] = field(default_factory=list)  # текст только содержательных таблиц
    tables_content: list[str] = field(default_factory=list)
    comments: list[str] = field(default_factory=list)
    pages_estimate: int = 0
    word_count: int = 0      # слов в параграфах (для отладки)
    is_appendix: bool = False
    in_doc_appendices: int = 0  # Приложения, встроенные в тело документа («ПРИЛОЖЕНИЕ №1»)


@dataclass
class AnalysisResult:
    """Analysis result for one document / group."""
    doc_name: str
    variants_count: int = 1
    pages: int = 1
    appendices: int = 0
    variables: int = 0
    variative_blocks: int = 0
    calculated_fields: int = 0
    tables: int = 0
    complexity: str = "Средняя"   # Высокая / Средняя / Низкая
    doc_type: str = ""             # тип договора (из заголовка)
    description: str = ""
    confidence: float = 0.0       # 0..1
    raw_llm: dict = field(default_factory=dict)
    found_variables: list = field(default_factory=list)   # перечень найденных переменных
    found_blocks: list = field(default_factory=list)      # перечень найденных вариативных блоков
    found_tables: list = field(default_factory=list)      # перечень найденных таблиц
    found_calculated: list = field(default_factory=list)  # перечень расчётных полей


# ─────────────────────────────────────────────
# STEP 1 — Extract content from .docx
# ─────────────────────────────────────────────

def _annotate_para(para) -> str:
    """
    Extract paragraph text with formatting annotations so the LLM can see
    color/highlight markers that indicate variable placeholders.

    Output tags injected into text:
      [МАРКЕР:текст]  — highlighted background (yellow/green/cyan) = variable placeholder
      [ЦВЕТ:текст]    — colored font (non-black) = variable OR instruction (model decides by context)
    """
    parts = []
    for run in para.runs:
        t = run.text
        if not t:
            continue

        is_highlighted = False
        is_colored = False
        try:
            if run.font.highlight_color is not None:
                is_highlighted = True
        except Exception:
            pass
        try:
            if run.font.color and run.font.color.type is not None:
                is_colored = True
        except Exception:
            pass

        # Только выделение фоном (highlight) = однозначно переменная.
        # Цветной шрифт НЕ аннотируем — включает инструкции разработчику (красный/синий).
        if is_highlighted and t.strip():
            parts.append(f"[МАРКЕР:{t.strip()}]")
        else:
            parts.append(t)

    return "".join(parts).strip()


def extract_docx(path: str) -> DocInfo:
    """Extract text, tables and comments from a .docx file."""
    doc = Document(path)
    name = Path(path).stem
    info = DocInfo(path=path, name=name)

    # Detect if it is an appendix (приложение) BY FILENAME
    info.is_appendix = bool(re.search(r"приложени", name, re.IGNORECASE))

    # ── Iterate document body children in ORDER so we can detect in-document
    # appendix sections and exclude their content from main body counts. ───────
    # doc.element.body direct children = top-level paragraphs (w:p) and tables
    # (w:tbl) only — cell-level paragraphs are nested deeper and not seen here.
    # This is the ONLY way to know "table T comes after heading ПРИЛОЖЕНИЕ №2".
    from docx.text.paragraph import Paragraph as _DocxPara
    from docx.table import Table as _DocxTable

    in_appendix_section = False   # флаг: прошли заголовок «ПРИЛОЖЕНИЕ №N»
    in_doc_app_count = 0

    for child in doc.element.body:
        child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if child_tag == 'p':
            para = _DocxPara(child, doc)
            plain = para.text.strip()

            # Если встретили заголовок-раздел «ПРИЛОЖЕНИЕ №1» — входим в режим приложения.
            # Фильтруем ложные срабатывания:
            #   • Список-перечень приложений в теле договора:
            #     «Приложение №1 – Перечень услуг;» → содержит тире-разделитель ' - '/' – '
            #   • Ссылочные предложения:
            #     «Приложение №1 является неотъемлемой частью...» → глагол в тексте
            #   • Слишком длинные строки (> 60 симв.) — скорее ссылка, чем автономный заголовок
            _APP_REF_VERBS = re.compile(
                r'\b(?:является|являются|прилагается|входит|считается|включает|'
                r'содержит|составляет|представляет|неотъемлемой|подписывается)\b',
                re.IGNORECASE
            )
            if (plain
                    and _APPENDIX_HEADING_RE.match(plain)
                    and len(plain) < 60
                    and not re.search(r'[\-–—]\s+\S', plain)      # нет «- описание»
                    and not _APP_REF_VERBS.search(plain)):          # нет ссылочных глаголов
                in_appendix_section = True
                in_doc_app_count += 1

            # Основная часть → info.paragraphs; приложения → info.appendix_paragraphs
            text = _annotate_para(para)
            if text:
                if not in_appendix_section:
                    info.paragraphs.append(text)
                else:
                    info.appendix_paragraphs.append(text)

        elif child_tag == 'tbl':
            info.tables_count += 1
            table = _DocxTable(child, doc)
            # ВАЖНО: python-docx возвращает объединённые ячейки несколько раз.
            # Используем id(cell._element) для дедупликации.
            rows_text = []
            seen_cell_ids: set = set()
            for row in table.rows:
                unique_cells = []
                for cell in row.cells:
                    cid = id(cell._element)
                    if cid not in seen_cell_ids:
                        seen_cell_ids.add(cid)
                        t = cell.text.strip()
                        if t:
                            unique_cells.append(t)
                if unique_cells:
                    rows_text.append(" | ".join(unique_cells))

            if rows_text:
                table_str = "\n".join(rows_text)
                is_service = bool(_SERVICE_TABLE_RE.search(table_str))
                if not in_appendix_section:
                    # Только основное тело: таблицы из приложений не считаем
                    info.tables_content.append(table_str)
                    if not is_service:
                        info.content_tables_count += 1
                        info.content_tables_content.append(table_str)

    info.in_doc_appendices = in_doc_app_count

    # Comments (stored in document.part XML)
    try:
        comments_part = doc.part.package.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
        )
        import xml.etree.ElementTree as ET
        tree = ET.fromstring(comments_part.blob)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for comment in tree.findall(".//w:comment", ns):
            parts = comment.findall(".//w:t", ns)
            text = " ".join(p.text or "" for p in parts).strip()
            if text:
                info.comments.append(text)
    except Exception:
        pass  # No comments part — that's fine

    # Страницы: читаем из XML + оцениваем по словам, берём максимум
    # (XML может быть устаревшим если документ не пересохранялся после правок)
    page_count_xml = None
    try:
        import zipfile as _zf_mod
        import xml.etree.ElementTree as _ET_mod
        with _zf_mod.ZipFile(path, 'r') as _zf:
            if 'docProps/app.xml' in _zf.namelist():
                with _zf.open('docProps/app.xml') as _f:
                    _root = _ET_mod.parse(_f).getroot()
                    for _elem in _root.iter():
                        if _elem.tag.split('}')[-1] == 'Pages' and _elem.text:
                            page_count_xml = max(1, int(_elem.text))
                            break
    except Exception:
        pass
    # Оцениваем страницы по словам из ВСЕГО документа (включая приложения и
    # таблицы — иначе многостраничные приложения занижают оценку).
    # doc.paragraphs — рекурсивный обход всего XML, включает ячейки таблиц.
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    info.word_count = word_count
    # ── Метод 1 (самый точный): lastRenderedPageBreak ──────────────────────────
    # Word вставляет эти элементы при каждом рендере — отражают РЕАЛЬНУЮ разбивку
    # на страницы с учётом шрифта, полей, интервалов. Работает если документ
    # был открыт и сохранён в Word хотя бы один раз.
    page_count_rendered = None
    try:
        _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        _lrpb = doc._element.findall(f'.//{{{_W_NS}}}lastRenderedPageBreak')
        if _lrpb:
            page_count_rendered = len(_lrpb) + 1   # разрывов + 1 = страниц
    except Exception:
        pass

    # ── Метод 2 (резервный): word_count / 400 ──────────────────────────────────
    # 400 слов/стр — умеренно плотный юридический текст.
    # Менее точен: не учитывает шрифт и межстрочный интервал.
    word_estimate = max(1, round(word_count / 400))

    # ── Выбираем наиболее надёжный источник ────────────────────────────────────
    if page_count_rendered is not None:
        # lastRenderedPageBreak — самый надёжный
        info.pages_estimate = page_count_rendered
    elif page_count_xml is not None:
        # XML (docProps/app.xml) — точен если документ пересохранялся,
        # но может быть устаревшим если страницы удалялись без пересохранения.
        # Сравниваем с word_estimate: если XML > word_estimate * 1.4 — подозрительно.
        ratio = page_count_xml / word_estimate if word_estimate > 0 else 99
        info.pages_estimate = page_count_xml if ratio <= 1.4 else word_estimate
    else:
        info.pages_estimate = word_estimate

    return info


# ─────────────────────────────────────────────
# STEP 1b — Deterministic variable counting
# ─────────────────────────────────────────────

def count_variables_deterministic(
    paragraphs: list[str],
    tables_content: list[str],
) -> tuple[int, list[str]]:
    """
    Count variables without LLM — no hallucination possible.

    Two sources:
      1. [МАРКЕР:text] tags — each unique normalised text = 1 variable.
      2. Underscore GROUPS: consecutive ___ runs connected by ≤25 non-letter chars
         are treated as ONE group (e.g. «___» _______ 20__ = one date variable).
         Each group is deduplicated by its word-based context key (last 3 meaningful
         words before the group start, skipping any truncated first word).

    Result is a conservative UPPER BOUND — the same variable may repeat in
    different sections, so actual unique count is usually somewhat lower.

    Returns: (total_count, debug_lines)
    """
    # Groups: consecutive ___ connected by non-letter gaps (dates, mixed placeholders)
    _VAR_GROUP_RE = re.compile(
        r'_{3,}'
        r'(?:[^a-zA-Zа-яёА-ЯЁ]{0,25}_{3,})*'
    )

    all_texts = paragraphs + tables_content

    # ── 1. [МАРКЕР:text] highlighted variables ──────────────────────────────
    # Две ситуации:
    #   a) highlight на текстовой метке («ФИО», «г._____») → named variable
    #   b) highlight на самом слоте (`___________`) → unnamed placeholder
    # Фильтрация мусора:
    #   • Чисто-подчёркивательные маркеры разной длины — это один тип слота.
    #   • Слишком длинные строки (> 80 симв.) — выделенные цветом условия/определения.
    #   • Выделенные тексты БЕЗ ___ и БЕЗ ключевых слов переменной — это определения
    #     терминов или названия сторон, подсвеченные для пояснения, НЕ переменные.
    #     Пример: «Вайлдберриз», «Группа компаний» — выделены как имена сторон,
    #     не являются пустыми слотами для заполнения.
    _PURE_UNDERSCORES_RE = re.compile(r'^[_\s\.,;()\[\]«»\'"–—№#\d]+$')
    # Слова-признаки переменной: если маркер не содержит ___ но содержит одно из
    # этих слов — это метка поля (label), тоже считается переменной.
    _VAR_LABEL_WORDS = frozenset({
        'фио', 'инн', 'кпп', 'огрн', 'огрнип', 'дата', 'номер', 'адрес',
        'наименование', 'должность', 'сумма', 'место', 'телефон', 'тел',
        'счёт', 'счет', 'бик', 'почт', 'email', 'mail', 'город', 'индекс',
        'паспорт', 'серия', 'подпись', 'период', 'срок', 'цена', 'стоимость',
    })
    markers: set[str] = set()
    for text in all_texts:
        for m in _MARKER_RE.finditer(text):
            val = m.group(1).strip().lower()
            if not val:
                continue
            # Слишком длинный → выделенное условие/определение, не переменная
            if len(val) > 80:
                continue
            # Состоит только из подчёркиваний, пунктуации, цифр → безымянный слот
            if _PURE_UNDERSCORES_RE.match(val):
                continue
            # Не содержит ___ — проверяем, есть ли признак переменной-метки.
            # Без ___ и без ключевого слова → скорее всего выделенный термин/название,
            # не переменная-слот (например «Вайлдберриз», «Группа компаний»).
            has_blank = bool(re.search(r'_{3,}', val))
            if not has_blank:
                val_words = set(re.findall(r'[а-яёa-z]{3,}', val))
                if not (val_words & _VAR_LABEL_WORDS):
                    continue
            markers.add(val)

    # ── 2. Underscore groups with word-based context deduplication ───────────
    under_keys: set[str] = set()
    unnamed_idx = 0
    # Pattern to find the party role label at the end of a preamble paragraph:
    # «именуемое(-ая/-ый) в дальнейшем «Поставщик»» / «далее – Заказчик»
    _ROLE_LABEL_RE = re.compile(
        r'(?:именуем\w+\s+в\s+дальнейшем|далее\s*[–—-])\s*[«"]?'
        r'(заказчик|покупатель|поставщик|исполнитель|подрядчик|арендодатель|'
        r'арендатор|продавец|лицензиар|лицензиат|займодавец|залогодатель|залогодержатель)',
        re.IGNORECASE
    )
    for text in all_texts:
        # Remove [МАРКЕР:...] wrappers to get plain text for context extraction
        plain = _MARKER_RE.sub(lambda mm: mm.group(1), text)

        # Detect paragraph-level party role (appears near end of preamble paragraphs)
        # Used to disambiguate structurally identical blanks for different parties.
        role_match = _ROLE_LABEL_RE.search(plain)
        para_role = _stem(role_match.group(1).lower()) if role_match else None

        for match in _VAR_GROUP_RE.finditer(plain):
            start = match.start()
            # Take 60 chars before the group start
            prefix_start = max(0, start - 60)
            prefix_raw = plain[prefix_start:start]

            # Extract all letter+digit tokens from the prefix
            words = re.findall(r'[a-zA-Zа-яёА-ЯЁ][a-zA-Zа-яёА-ЯЁ0-9]*', prefix_raw.lower())

            # Drop truncated fragment ONLY if the window was cut in the middle of a word
            # (i.e., the window start is not at position 0 AND prefix starts with a letter)
            if (prefix_start > 0
                    and prefix_raw
                    and prefix_raw[0] not in ' \t\n.,;:!?«»"\'()–—/\\'):
                words = words[1:]

            # Drop very short tokens (1-2 char articles, prepositions, noise)
            # + нормализуем падежи стеммером чтобы «заказчика»/«заказчику» → один ключ
            words = [_stem(w) for w in words if len(w) > 2]

            if words:
                # Key = last 3 stemmed words before the placeholder.
                # Prepend party role if detected so that «в лице Поставщика» and
                # «в лице Покупателя» get distinct keys instead of both being «лиц».
                key_parts = words[-3:]
                if para_role:
                    key_parts = [para_role] + key_parts
                key = ' '.join(key_parts)
                under_keys.add(key)
            else:
                # No usable context — count as distinct unnamed placeholder
                unnamed_idx += 1
                under_keys.add(f"<unnamed_{unnamed_idx}>")

    # ── 3. Semantic merge: collapse keys sharing 2+ common words ─────────────
    # Same variable can appear at different positions → different 3-word keys.
    # Merging them: «договору подрядчик является» + «настоящему договору подрядчик» → 1.
    key_list = sorted(under_keys)
    key_words_list = [frozenset(k.split()) for k in key_list]
    merged: set[str] = set()
    used = [False] * len(key_list)
    for i in range(len(key_list)):
        if used[i]:
            continue
        merged.add(key_list[i])
        wi = key_words_list[i]
        for j in range(i + 1, len(key_list)):
            if not used[j] and len(wi & key_words_list[j]) >= 2:
                used[j] = True
    under_keys = merged

    # ── 4. Raw occurrence count (before deduplication) ──────────────────────
    # Count total physical underscore groups across all texts (with repetitions).
    # This gives a rough ceiling: the same field can appear in multiple places,
    # so raw_count >= true unique variable count.
    raw_under_count = 0
    for text in all_texts:
        plain = _MARKER_RE.sub(lambda mm: mm.group(1), text)
        raw_under_count += len(list(_VAR_GROUP_RE.finditer(plain)))

    # ── 5. Build debug output ───────────────────────────────────────────────
    debug: list[str] = []
    debug.append(f"  Highlights        : {len(markers)} unique values")
    for v in sorted(markers):
        debug.append(f"    [МАРКЕР] «{v}»")
    debug.append(f"  Underscore groups : {len(under_keys)} unique contexts (raw={raw_under_count})")
    for k in sorted(under_keys):
        debug.append(f"    [___] key='{k}'")

    # ── 6. Выбор источника ──────────────────────────────────────────────────
    if markers or under_keys:
        if len(markers) >= len(under_keys):
            total = len(markers)
            source = "highlight markers (larger)"
        else:
            total = len(under_keys)
            source = "underscore groups (larger)"
    else:
        total = 0
        source = "none found"
    debug.insert(0, f"  Python var count: {total} deduped / {raw_under_count} raw"
                    f"  (markers={len(markers)}, underscore_groups={len(under_keys)}, source={source})")
    return total, raw_under_count, debug


# ─────────────────────────────────────────────
# STEP 1c — Deterministic variative block counting
# ─────────────────────────────────────────────

def count_variative_python(paragraphs: list[str]) -> tuple[int, list[str]]:
    """
    Count variative block markers in ALL paragraphs.
    Sources:
      1. [условие/инструкция] bracket markers (explicit bracket conditions)
      2. «Вариант №1 / Вариант А» paragraph starters — КАЖДЫЙ стартовый маркер
         (Вариант 1, Вариант А) означает отдельный блок альтернатив.
    Lower bound: may miss implicit conditional sentences, Word comments, etc.

    Returns: (count, debug_lines)
    """
    # Длина 3-300: русские инструкции разработчику часто длиннее 100 символов.
    # Например: «[Следующие условия применяются только для варианта с рабочей документацией]» = 80 симв.
    # Увеличиваем с 100 до 300 чтобы не пропускать длинные инструкции.
    _BRACKET_RE = re.compile(r'\[([^\]]{3,300})\]')
    # Paragraph starts with «Вариант №1», «Вариант 1:», «Вариант А» etc.
    # Ловим ТОЛЬКО первый вариант группы (1, А) — не «Вариант 2», «Вариант Б»
    # чтобы не дублировать: «Вариант 1 + Вариант 2» = 1 блок, не 2.
    _VARIANT_START_RE = re.compile(
        r'^вариант\s*(?:№|#)?\s*(?:1|i|а|a)(?:\s|[:.()\-]|$)',
        re.IGNORECASE
    )
    # Words indicating a field name — skip these (they are variables, not blocks)
    FIELD_WORDS = frozenset({
        'фио', 'инн', 'кпп', 'огрн', 'огрнип', 'дата', 'номер', 'адрес',
        'наименование', 'должность', 'сумма', 'маркер', 'место', 'телефон',
        'счёт', 'бик', 'описание', 'реквизит',
    })

    bracket_found: set[str] = set()
    variant_found: set[str] = set()
    debug_items: list[str] = []

    for text in paragraphs:
        # ── 1. [brackets] ────────────────────────────────────────────────────
        for m in _BRACKET_RE.finditer(text):
            content = m.group(1).strip()
            content_lower = content.lower()
            if content_lower.startswith('маркер:'):
                continue
            # Если скобки содержат подчёркивания — это переменная, не условие
            if re.search(r'_{3,}', content):
                continue
            words = set(re.findall(r'[а-яёa-z]+', content_lower))
            if words & FIELD_WORDS:
                continue
            key = content_lower[:60]
            if key not in bracket_found:
                bracket_found.add(key)
                debug_items.append(f"    [БЛОК] «{content[:70]}»")

        # ── 2. «Вариант №1 / Вариант А» paragraph starters ──────────────────
        # Убираем [МАРКЕР:...] чтобы не сбивало match
        plain = _MARKER_RE.sub(lambda mm: mm.group(1), text).strip()
        if _VARIANT_START_RE.match(plain):
            key = plain[:80].lower()
            if key not in variant_found:
                variant_found.add(key)
                debug_items.append(f"    [ВАРИАНТ] «{plain[:70]}»")

    total = len(bracket_found) + len(variant_found)
    debug: list[str] = [
        f"  Python variative count (lower bound): {total} "
        f"(brackets={len(bracket_found)}, variant_headers={len(variant_found)})"
    ] + debug_items
    return total, debug


# ─────────────────────────────────────────────
# STEP 2 — Cross-document comparison
# ─────────────────────────────────────────────

def diff_paragraph_sets(texts_a: list[str], texts_b: list[str]) -> dict:
    """
    Compare two lists of paragraphs and return:
    - only_in_a  : paragraphs unique to document A
    - only_in_b  : paragraphs unique to document B
    - changed    : paragraphs that look similar but differ (Levenshtein-light)
    """
    set_a = set(texts_a)
    set_b = set(texts_b)
    only_a = [t for t in texts_a if t not in set_b]
    only_b = [t for t in texts_b if t not in set_a]

    # Simple "near-match" detection — paragraphs that share a long prefix
    changed = []
    for pa in only_a:
        for pb in only_b:
            prefix_len = _common_prefix_len(pa, pb)
            if prefix_len > 40 and pa != pb:
                changed.append({"a": pa, "b": pb})
                break

    return {
        "only_in_a": only_a,
        "only_in_b": only_b,
        "changed": changed,
    }


def _common_prefix_len(a: str, b: str) -> int:
    n = min(len(a), len(b))
    for i in range(n):
        if a[i] != b[i]:
            return i
    return n


def build_diff_summary(pairs: list[tuple[DocInfo, DocInfo]]) -> str:
    """
    Build a text summary of differences across all provided pairs.
    This goes into the LLM prompt as context.
    """
    lines = []
    for doc_a, doc_b in pairs:
        diff = diff_paragraph_sets(doc_a.paragraphs, doc_b.paragraphs)
        lines.append(f"\n=== Сравнение: «{doc_a.name}» vs «{doc_b.name}» ===")
        lines.append(f"Уникальных параграфов в A: {len(diff['only_in_a'])}")
        lines.append(f"Уникальных параграфов в B: {len(diff['only_in_b'])}")
        lines.append(f"Изменённых параграфов: {len(diff['changed'])}")

        if diff["only_in_a"]:
            lines.append("\nТолько в A (первые 5):")
            for p in diff["only_in_a"][:5]:
                lines.append(f"  - {p[:120]}")

        if diff["only_in_b"]:
            lines.append("\nТолько в B (первые 5):")
            for p in diff["only_in_b"][:5]:
                lines.append(f"  - {p[:120]}")

        if diff["changed"]:
            lines.append("\nИзменённые параграфы (первые 3):")
            for c in diff["changed"][:3]:
                lines.append(f"  A: {c['a'][:100]}")
                lines.append(f"  B: {c['b'][:100]}")
                lines.append("")

    return "\n".join(lines)


# ─────────────────────────────────────────────
# STEP 3 — Claude API (few-shot) analysis
# ─────────────────────────────────────────────

FEW_SHOT_SYSTEM = """Ты — аналитик документов для платформы TurboContract (конструктор юридических договоров).
Твоя задача — проанализировать юридический документ и вернуть JSON с параметрами для расчёта трудозатрат разработчика шаблона.

━━━ ОПРЕДЕЛЕНИЯ ━━━

АННОТАЦИИ ФОРМАТИРОВАНИЯ: система предобработки автоматически размечает отформатированный текст тегами:
  [МАРКЕР:текст] — текст с цветовым выделением фона (жёлтый/зелёный/голубой highlight) = ПЕРЕМЕННАЯ
  [ЦВЕТ:текст]   — текст с цветным шрифтом (красный/синий/серый и т.п.) = ПЕРЕМЕННАЯ или ИНСТРУКЦИЯ
                   (ЦВЕТ = инструкция разработчику если это пояснение/примечание; ЦВЕТ = переменная если это незаполненное поле)

ПЕРЕМЕННАЯ (variables) — уникальное место, куда подставляется конкретное значение при генерации документа.
Маркеры переменных в реальных документах:
  • [МАРКЕР:текст] — переменная ТОЛЬКО если содержимое является ПУСТЫМ БЛАНКОМ или МЕТКОЙ ПОЛЯ.
    Три типа [МАРКЕР:] — их надо различать:
    ✓ БЛАНК: [МАРКЕР:_________] или [МАРКЕР:___] — пустое место для заполнения = переменная
    ✓ МЕТКА ПОЛЯ: [МАРКЕР:ФИО], [МАРКЕР:ИНН], [МАРКЕР:г.__] — метка типа переменной = переменная
    ⛔ ГОТОВЫЙ ТЕКСТ: [МАРКЕР:Вайлдберриз], [МАРКЕР:ООО «РВБ»] — реальное название/термин = НЕ переменная!
       (выделен для акцента, не для заполнения)
    ⛔ УСЛОВИЕ/ПРЕДЛОЖЕНИЕ: [МАРКЕР:Обязательные условия — условия при заключении договоров с ВБ...]
       — целое предложение/абзац = ВАРИАТИВНЫЙ БЛОК, НЕ переменная!
    ⚠ НЕСКОЛЬКО [МАРКЕР:] ПОДРЯД = часто ОДНО поле — считай как одну переменную:
       Дата: [МАРКЕР:№ ][МАРКЕР:__ ][МАРКЕР:от ][МАРКЕР:_][МАРКЕР:.][МАРКЕР:.20][МАРКЕР:__] г. → 1 переменная
       Сумма: [МАРКЕР:_____][МАРКЕР:(][МАРКЕР:_____][МАРКЕР:)] рублей → 1 переменная
    ⚠ [МАРКЕР:] содержащий ТОЛЬКО пунктуацию/пробелы («», (), ., ,) — соединитель, НЕ переменная.
  • [ЦВЕТ:текст] — переменная, если содержимое похоже на незаполненное поле (имя, дата, сумма, адрес)
  • _____________ (подчёркивания любой длины) — самый распространённый текстовый маркер
  • «__» _______ 20__ г. / "___" _______ — формат даты с подчёркиваниями
  • № _____ / No. _____ — номер договора
  • ____ (___) дней — числовое значение + прописью в скобках
  • [ название поля ] — квадратные скобки = ПЕРЕМЕННАЯ только если содержимое является
    НАЗВАНИЕМ ПОЛЯ для заполнения: [ФИО], [дата], [ИНН], [номер], [адрес], [наименование], [должность] и т.п.
  • ⛔ [ условие или инструкция ] — НЕ переменная, это МАРКЕР ВАРИАТИВНОГО БЛОКА (см. ниже).
    Работает для ЛЮБОЙ длины — даже короткие условия:
    [если ЮЛ], [для ИП], [удалить], [при НДС], [оставить] = вариативный блок, НЕ переменная.
  • ПЕРЕМЕННАЯ-СЕЛЕКТОР — переменная, содержащая перечень взаимоисключающих вариантов через
    «/» или «или»: [ЮЛ / ИП / ФЛ], «тип договора: аванс/постоплата», «___ (с НДС / без НДС)».
    Такая переменная УПРАВЛЯЕТ вариативными блоками в других местах документа — см. ниже.
  • (  ) / ( текст ) — круглые скобки как заполнители
  • Пустое место после двоеточия: «Дата:      »

ТИПОВЫЕ СИСТЕМНЫЕ ПЕРЕМЕННЫЕ TurboContract (всегда есть в любом договоре — ищи их в первую очередь):
  Документ: Номер, Наименование, Дата, Место составления, Тип договора
  Наше лицо: Тип лица, Признак ИП, ОГРНИП, Дата записи ЕГРИП, ФИО (если ФЛ),
             Полное наименование, Сокращённое наименование, ОГРН, ИНН, КПП,
             Юридический адрес, ФИО подписанта, Должность подписанта, Документ-основание полномочий
  Контрагент: те же 14 полей что у «нашего лица» (тип лица, ИП/ЮЛ/ФЛ, реквизиты, подписант)
  Финансы: Признак НДС, Ставка НДС, Сумма с НДС, Вычисление суммы НДС, Сумма без НДС
  ИТОГО типовой блок: ~33 системные переменные (5 + 14 + 14 + 5)

ВАЖНО — правила счёта:
  • Считай УНИКАЛЬНЫЕ переменные, не вхождения. Если «ФИО подписанта» встречается 5 раз — это 1 переменная.
  • ⛔ НЕ добавляй переменные автоматически. Считай ТОЛЬКО те поля, которые РЕАЛЬНО присутствуют в тексте как незаполненные маркеры (___, [], цветовое выделение и т.д.). Если в документе нет поля ИНН — не добавляй его.
  • Список из 33 типовых системных переменных дан для РАСПОЗНАВАНИЯ: увидел ___ рядом с «ИНН», «ОГРН», «ФИО» — это системная переменная, считай. Не увидел в тексте — не добавляй.
  • Подписи в конце документа НЕ считаются отдельными переменными — они повторяют реквизиты.
  • Инструкции разработчику (красный/синий курсив, служебные комментарии) — НЕ переменные.
  • ДВУЯЗЫЧНЫЕ документы (RUS+ENG): если один и тот же документ содержит параллельный перевод, считай переменные ТОЛЬКО ОДИН РАЗ. «Дата ___» и «Date ___» = 1 переменная, не 2.
  • ⛔ ПРИЛОЖЕНИЯ: бланки в приложениях (Приложение №1, №2 и т.д.) — НЕ считай как переменные шаблона!
    Поля в приложениях заполняются вручную при исполнении договора, а не через шаблонизатор TurboContract.
    Переменные считай ТОЛЬКО из основного тела договора (до первого «ПРИЛОЖЕНИЕ №»).

ВАРИАТИВНЫЙ_БЛОК (variative_blocks) — пункт/раздел/абзац, который целиком включается ИЛИ исключается в зависимости от условия. Также: альтернативные формулировки одного пункта (выбрать одну из нескольких).

Маркеры вариативности — ищи ВСЕ эти признаки:

  ① [Условие или инструкция в квадратных скобках] — САМЫЙ ЧАСТЫЙ МАРКЕР!
     Главное правило: [ ] = вариативный блок, если содержимое — УСЛОВИЕ или ИНСТРУКЦИЯ,
     а не название поля. Работает для ЛЮБОЙ длины текста.
     Короткие: [если ЮЛ], [для ИП], [удалить], [при НДС], [оставить] = 1 блок каждый.
     Длинные: [если Покупатель — юридическое лицо, использовать этот пункт] = 1 блок.
     Ключевой вопрос: «это ИНСТРУКЦИЯ что делать или НАЗВАНИЕ ПОЛЯ для заполнения?»
     → Инструкция = вариативный блок. Название поля = переменная.

  ② ПЕРЕМЕННАЯ-СЕЛЕКТОР создаёт вариативные блоки — самый скрытый тип вариативности!
     Признак: переменная с вариантами выбора [ЮЛ/ИП/ФЛ], «тип: A или B», «___ (вариант 1/2)».
     Как считать:
       a) Найди переменную-селектор в документе
       b) Найди все блоки/разделы, которые зависят от её значения (применяются только для ЮЛ,
          только для ИП и т.д.) — каждый такой блок = 1 вариативный блок
       c) Если один селектор управляет N блоками в разных местах — считай все N
     Пример: «Тип покупателя [ЮЛ / ИП / ФЛ]» → управляет блоком реквизитов (3 варианта),
     блоком НДС (2 варианта), блоком ответственности (2 варианта) = итого ~7 вариативных блоков
     от одной переменной-селектора.

  ③ НЕСКОЛЬКО ВЕРСИЙ ДОКУМЕНТА (cross-document diff) — если клиент прислал несколько
     похожих документов для «склейки»: каждый раздел/абзац, присутствующий в одних версиях
     и отсутствующий в других = 1 вариативный блок. Смотри раздел «Результат сравнения версий»
     в описании документа — он уже содержит подсчёт таких различий.

  ④ Word-комментарии (аннотации справа в документе) — если комментарий описывает условие
     применения абзаца или альтернативный вариант текста = 1 вариативный блок за каждый
     смысловой комментарий о вариативности.

  ⑤ «Вариант №1 / Вариант №2 / Вариант №3» — явный маркер альтернатив
  ⑥ «Вариант А / Вариант Б», «вариант 1: ... вариант 2: ...» — альтернативные формулировки
  ⑦ «Выбрать вариант», «Выбрать, если применимо» — инструкция выбора
  ⑧ «ненужное зачеркнуть» / «имеет/не имеет» / «да/нет» — выбор из двух значений
  ⑨ «в случае если» / «при наличии» / «если применимо» — условный блок ТОЛЬКО если ВЕСЬ ПУНКТ
     является опциональным и помечен явным маркером (МАРКЕР/[], Вариант А/Б, инструкция).
     ⛔ НЕ считай условные предложения ВНУТРИ обязательного пункта!
     Пример НЕ блока: «Заказчик вправе удержать штраф в случае просрочки» — это фиксированное условие.
     Пример БЛОКА: «[если применяется НДС] следующий пункт применяется» — весь пункт опционален.
  ⑩ Курсивный или цветной текст с альтернативной формулировкой внутри основного абзаца

ВАЖНО — правила счёта вариативных блоков:
  • Каждая пара «Вариант А / Вариант Б» = 1 блок (не 2). Каждый [ ] с условием = 1 блок.
  • ⛔ НЕ считай как вариативные блоки: стандартные условия договора без явных маркеров выбора,
    обычные условные конструкции внутри пунктов («при просрочке», «в случае нарушения» и т.п.).
  • ОДИН МАРКЕР [условие] в начале РАЗДЕЛА из нескольких абзацев = СТОЛЬКО блоков, СКОЛЬКО абзацев
    в этом разделе (каждый абзац, к которому применяется условие — отдельный блок для шаблона).
    Пример: [только для ПД+РД] → 5 абзацев под условием → 5 вариативных блоков (НЕ 1).
  • Аналогично: «Следующие N абзацев относятся к варианту X» = N блоков.
  • ТОЛЬКО явные маркеры суммируются — НЕ угадывай скрытую вариативность.

РАСЧЁТНОЕ_ПОЛЕ (calculated_fields) — значение, которое TurboContract АВТОМАТИЧЕСКИ вычисляет при генерации (формульный движок платформы).
ЯВЛЯЕТСЯ расчётным ТОЛЬКО если выполнены ОБА условия:
  (1) в документе есть ТАБЛИЦА-спецификация с числовыми данными (кол-во, цена и т.п.), И
  (2) итоговая ячейка таблицы вычисляется по формуле (итог = кол-во × цену, сумма строк, автосумма НДС).
⛔ НЕ является расчётным (по умолчанию 0 если нет таблицы-спецификации):
  • пени/неустойки «0,1% в сутки» — текстовая ставка, НЕ авторасчёт
  • гарантийное удержание «N% от цены» — текстовая формула, НЕ авторасчёт
  • НДС «ставка X%» когда сумма просто упомянута в тексте — НЕ авторасчёт
  • процентные ставки, штрафы «M% от суммы» — это переменные или текст
  • любые суммы просто указанные в тексте договора
ПРАВИЛО ПО УМОЛЧАНИЮ: если сомневаешься — ставь 0. Расчётные поля редки (только крупные договоры со спецификацией-таблицей).

ТАБЛИЦА (tables) — считай все таблицы с предметными данными договора в ОСНОВНОМ документе.
Считай (примеры):
  • Спецификации товаров/услуг (перечень с количеством, ценой и т.п.)
  • Расчёты, сметы, графики платежей
  • Перечни работ, этапов, разделов документации (например «№ п/п | Раздел | Ссылка»)
  • Перечни требований, нормативов, объектов
  • Таблицы условий, сроков, штрафных шкал
НЕ считай:
  • Таблицы реквизитов и подписей сторон в конце документа (банковские реквизиты, адреса, блоки подписей «ЗАКАЗЧИК | ПОДРЯДЧИК»)
  • Одиночные ячейки-заголовки (просто название компании)
  • Таблицы из приложений

ПРИЛОЖЕНИЕ (appendices) — отдельный документ, на который есть ссылка «Приложение №N».

СЛОЖНОСТЬ (считай на основе реально найденных переменных и блоков):
  • Высокая  — >60 переменных ИЛИ >35 вариативных блоков ИЛИ есть расчётные поля
  • Средняя  — 30–60 переменных ИЛИ 10–35 вариативных блоков
  • Низкая   — <30 переменных И <10 вариативных блоков, расчётных полей нет

━━━ ФОРМАТ ОТВЕТА ━━━
Верни ТОЛЬКО валидный JSON без комментариев и без markdown-обёрток (без ```):
{
  "variables": <int>,
  "variative_blocks": <int>,
  "calculated_fields": <int>,
  "tables": <int>,
  "complexity": "Высокая" | "Средняя" | "Низкая",
  "doc_type": "<тип договора полностью, напр. 'Договор аренды оборудования', 'Договор подряда на ПИР'>",
  "description": "<1–2 предложения: тип документа, стороны, суть вариативности>",
  "confidence": <float 0.0–1.0>,
  "notes": "<что сложно определить, требует проверки человеком>",
  "found_variables": ["<краткое название переменной 1>", "<переменная 2>", ...],
  "found_blocks": ["<краткое описание блока/условия 1>", "<блок 2>", ...],
  "found_tables": ["<описание таблицы 1>", "<таблица 2>", ...],
  "found_calculated": ["<описание расчётного поля 1>", ...]
}

ВАЖНО для doc_type:
• Извлекай из заголовка документа или преамбулы (первые 3–5 абзацев).
• Пиши полное название без сокращений: не "Договор ПИР", а "Договор на выполнение проектно-изыскательских работ".
• Если тип неоднозначен — укажи наиболее точный по содержанию.

ВАЖНО для found_variables / found_blocks / found_tables / found_calculated:
• Каждый элемент — короткое название (3–8 слов), не копируй весь абзац.
• ⛔ ЗАПРЕЩЕНО использовать аббревиатуры без расшифровки. Всегда пиши полностью:
  БГ → Банковская гарантия, ПНР → Пусконаладочные работы, РД → Рабочая документация,
  ПИР → Проектно-изыскательские работы, ТЗ → Техническое задание, КС → Корреспондентский счёт,
  НДС → Налог на добавленную стоимость, ЭВМ → Электронная вычислительная машина.
  Если аббревиатура раскрыта в тексте договора — используй то раскрытие.
• found_variables: перечисли каждую уникальную переменную (напр. "Номер договора", "Дата", "Сумма", "ФИО Заказчика").
• found_blocks: опиши условие каждого блока (напр. "Если предусмотрен аванс", "Вариант без авансового платежа").
• found_tables: опиши содержание таблицы (напр. "Спецификация товаров", "Таблица реквизитов сторон").
• found_calculated: если расчётных полей нет — верни пустой массив []."""


FEW_SHOT_EXAMPLES = [

    # ── Пример 1: Сублицензионный договор (Низкая, реальный документ) ──
    {
        "role": "user",
        "content": """Проанализируй документ:
--- ДОКУМЕНТ: ТФ_сублицензионного_договора ---
Сублицензионный договор № ______________
_______ «         », именуемое в дальнейшем «Лицензиат», в лице______________, действующего на основании __________, с одной стороны, и
Страховое публичное акционерное общество «Ингосстрах», именуемое в дальнейшем «Сублицензиат», в лице _________________, действующего на основании__________________, с другой стороны, заключили настоящий Договор о нижеследующем:

ПРЕДМЕТ ДОГОВОРА
Лицензиат обязуется предоставить Сублицензиату право использования программ для ЭВМ. Наименование программ, территория и срок использования указаны в Приложении №1 (Спецификация).
Исключительное право на ПО принадлежит ___________ (Правообладатель), а Лицензиат действует в пределах прав по ____________________ № ___________ от ________________г., заключённому между _______________ и ___________________.

РАЗМЕР ВОЗНАГРАЖДЕНИЯ
Вознаграждение по Договору составляет __________________________,
в течение 20 рабочих дней с даты подписания Акта приема-передачи ПО Сублицензиат перечисляет вознаграждение в размере 100% стоимости.
Если в Спецификации цена указана в валюте, платежи осуществляются в рублях РФ по курсу ЦБ РФ на день платежа.
--- КОНЕЦ ---
Количество вариантов документа: 1. Таблиц Word: 1. Приложений: 1 (Спецификация).""",
    },
    {
        "role": "assistant",
        "content": """{
  "variables": 12,
  "variative_blocks": 1,
  "calculated_fields": 0,
  "tables": 1,
  "complexity": "Низкая",
  "doc_type": "Сублицензионный договор на программное обеспечение",
  "description": "Сублицензионный договор на ПО между Лицензиатом и СПАО Ингосстрах. Один вариативный блок — оплата в рублях или валюте по курсу ЦБ.",
  "confidence": 0.88,
  "notes": "Переменные явные (___). Вариативность минимальная — только по валюте платежа.",
  "found_variables": ["Номер договора", "Название Лицензиата", "Представитель Лицензиата", "Основание Лицензиата", "Представитель Сублицензиата", "Основание Сублицензиата", "Название Правообладателя", "Номер лицензионного договора", "Дата лицензионного договора", "Сторона 1 лицензионного договора", "Сторона 2 лицензионного договора", "Размер вознаграждения"],
  "found_blocks": ["Оплата в рублях по курсу ЦБ (если цена в валюте)"],
  "found_tables": ["Спецификация программ (Приложение №1)"],
  "found_calculated": []
}""",
    },

    # ── Пример 2: Договор-конструктор ИТ (Низкая, реальный документ) ──
    {
        "role": "user",
        "content": """Проанализируй документ:
--- ДОКУМЕНТ: ТФ_договора-конструктора ---
Договор №
г. Москва   «___» _______ 202_ года
СПАО «ИНГОССТРАХ», в лице ______________________, действующего на основании __________________, с одной стороны, и
________ «_______», в лице _____________________, действующего на основании ______________, с другой стороны, заключили настоящий Договор.

Предмет Договора
Исполнитель обязуется выполнить Работы по выбрать вариант в любой комбинации из списка:
внедрению, разработке, доработке, установке, тестированию, сопровождению, адаптации, модификации программы для ЭВМ_____, программного продукта________, баз данных _______, системы ________.

Выбрать, если применимо
Вариант 1 (для рамочного договора): в соответствии со Спецификациями по форме Приложения №1.
Вариант 2 (для договора на конкретный объём работ): в соответствии с Приложением №1.

Вариант 1: Работы выполняются по месту нахождения Исполнителя (удалённо). Необходимость выезда согласуется с Заказчиком.
Вариант 2: Работы выполняются Исполнителем на территории Заказчика с использованием предоставленных материалов.

Интеллектуальные права
Вариант 1: Исключительные права в полном объёме переходят к Заказчику.
Вариант 2: Исполнитель обязуется выполнять Работы своими силами или привлечь _____ «____», ИНН ____________.

Срок выполнения Работ: _________________.
Стоимость Работ: _________________ рублей (в т.ч. НДС _____ / без НДС).
--- КОНЕЦ ---
Количество вариантов документа: 1 (варианты внутри документа). Таблиц Word: 1. Приложений: нет явных.""",
    },
    {
        "role": "assistant",
        "content": """{
  "variables": 21,
  "variative_blocks": 7,
  "calculated_fields": 0,
  "tables": 1,
  "complexity": "Низкая",
  "doc_type": "Договор на разработку и внедрение программного обеспечения",
  "description": "ИТ-договор на разработку/внедрение ПО (СПАО Ингосстрах). Вариативность по типу работ, месту выполнения, рамочный/конкретный объём, переходу прав, привлечению субподрядчика.",
  "confidence": 0.91,
  "notes": "Вариативные блоки явно помечены 'Вариант 1/2'. Переменные — ___ и пустые строки. Сложность Низкая несмотря на 7 блоков — сами блоки короткие.",
  "found_variables": ["Номер договора", "Дата", "Место заключения", "Название Исполнителя", "Представитель Исполнителя", "Основание Исполнителя", "Название Заказчика (СПАО Ингосстрах — фиксировано)", "Представитель Заказчика", "Основание Заказчика", "Наименование ПО/системы", "ИНН субподрядчика (если привлекается)", "Название субподрядчика", "Сумма договора", "Срок выполнения", "Банковские реквизиты Исполнителя", "Адрес Исполнителя", "Телефон/email Исполнителя", "Подпись/ФИО Исполнителя", "Подпись/ФИО Заказчика", "Печать Исполнителя", "Печать Заказчика"],
  "found_blocks": ["Вариант 1: рамочный договор со спецификациями", "Вариант 2: договор на конкретный объём работ", "Вариант 1: работы выполняются удалённо", "Вариант 2: работы выполняются на территории Заказчика", "Вариант 1: исключительные права переходят к Заказчику", "Вариант 2: привлечение субподрядчика", "Выбор применимого типа работ из списка"],
  "found_tables": ["Таблица реквизитов и подписей сторон"],
  "found_calculated": []
}""",
    },

    # ── Пример 3: Геолого-технологический договор (Высокая, реальный документ) ──
    {
        "role": "user",
        "content": """Проанализируй документ:
--- ДОКУМЕНТ: Договор_на_проведение_геолого_технологических_исследований ---
Договор №_______________
на проведение геолого-технологических исследований и газового каротажа в процессе бурения скважин

____________________________________________________ (__________), именуемое в дальнейшем Заказчик, в лице _________________________, действующего на основании _______, с одной стороны, и
______________________________________ (____________), именуемое в дальнейшем Подрядчик, в лице ________________________, действующего на основании _______________, с другой стороны, заключили настоящий Договор.

1.2. Подрядчик обязуется выполнять работы в соответствии с Техническим заданием (Приложение №1).
1.3. Дата начала выполнения работ согласовывается при подаче заявки. Время монтажа/демонтажа — не более 3 суток.

Все работы выполняются в соответствии с ПБНГП, Технической инструкции по проведению ГТИ нефтяных и газовых скважин [РД 153-39.0-069-01].

3.1. Стоимость работ по настоящему Договору составляет ____________ (__________________) рублей в сутки, в том числе НДС 20% - ______________ рублей.
3.2. Стоимость монтажа/демонтажа составляет ___________________ (____________________) рублей за операцию, в том числе НДС 20% - ___________ рублей.
3.3. Стоимость мобилизации/демобилизации составляет ______________ (_______________________) рублей за операцию.
3.4. В случае возникновения простоя по вине Заказчика, Заказчик оплачивает простой в размере _________ (___________________) рублей в сутки.
3.5. В случае возникновения простоя оборудования без персонала по вине Заказчика, стоимость простоя составляет _____________ (_______________) рублей в сутки.

7.2. За нарушение срока оплаты Заказчик уплачивает неустойку в размере 0,1% от неоплаченной суммы за каждый день просрочки.

Для направления информации Контрагенту: тел. ______________________, эл.адрес: ______________________.

РЕКВИЗИТЫ И ПОДПИСИ СТОРОН
Заказчик: ___________________ / ___________________
Подрядчик: __________________ / __________________
--- КОНЕЦ ---
Количество вариантов документа: 1. Таблиц Word: 6. Приложений: 8.
Стиль разметки: переменные выделены цветом, маркеры ___ и [ ], ссылки на НПА в [скобках] как вариативные блоки.""",
    },
    {
        "role": "assistant",
        "content": """{
  "variables": 50,
  "variative_blocks": 4,
  "calculated_fields": 0,
  "tables": 6,
  "complexity": "Высокая",
  "doc_type": "Договор на проведение геолого-технологических исследований и газового каротажа скважин",
  "description": "Договор на геолого-технологические исследования и газовый каротаж скважин. Переменные — реквизиты сторон, стоимость суточная/монтаж/мобилизация/простой. Вариативность через ссылки на НПА в [скобках] и условия 'в случае если'.",
  "confidence": 0.82,
  "notes": "Расчётных полей НЕТ: пени 0,1% в сутки — текстовая формула в договоре, не авторасчёт платформы TurboContract; НДС 20% — фиксированная ставка в тексте, не авторасчёт (нет таблицы-спецификации с переменными кол-во×цена). Вариативные блоки: 'в случае возникновения простоя' (×2 условия) + ссылки на [РД...]. Рекомендуется ручная проверка числа переменных ±8.",
  "found_variables": ["Номер договора", "Название Заказчика", "Краткое название Заказчика", "Представитель Заказчика", "Основание Заказчика", "Название Подрядчика", "Краткое название Подрядчика", "Представитель Подрядчика", "Основание Подрядчика", "Стоимость работ в сутки (цифры)", "Стоимость работ в сутки (прописью)", "Стоимость монтажа/демонтажа (цифры)", "Стоимость монтажа/демонтажа (прописью)", "Стоимость мобилизации (цифры)", "Стоимость мобилизации (прописью)", "Стоимость простоя с персоналом (цифры)", "Стоимость простоя с персоналом (прописью)", "Стоимость простоя без персонала (цифры)", "Стоимость простоя без персонала (прописью)", "Телефон контрагента", "Email контрагента", "Реквизиты Заказчика", "Реквизиты Подрядчика", "ФИО подписанта Заказчика", "ФИО подписанта Подрядчика"],
  "found_blocks": ["Ссылка на РД 153-39.0-069-01 (применимость НПА)", "Простой по вине Заказчика с персоналом", "Простой оборудования без персонала по вине Заказчика", "Условия привлечения субподрядчика"],
  "found_tables": ["Таблица реквизитов и подписей сторон", "Техническое задание (Приложение №1)", "Прочие приложения (8 шт.)"],
  "found_calculated": []
}""",
    },

    # ── Пример 4: Индивидуальные условия залога (Высокая, реальный документ) ──
    {
        "role": "user",
        "content": """Проанализируй документ:
--- ДОКУМЕНТ: Индивидуальные_условия_Договора_залога ---
(Текст и сноски курсивом синего цвета — инструкции, неиспользуемые варианты подлежат удалению.)

ИНДИВИДУАЛЬНЫЕ УСЛОВИЯ №      Договора залога
«  »       20   г.    г.

СТОРОНЫ ДОГОВОРА ЗАЛОГА
Залогодатель:
Вариант №1 (для юридических лиц)
(полное и краткое наименование)  ОГРН     ИНН
ФИО и наименование единоличного исполнительного органа
Адрес места нахождения:    Почтовый адрес:    Номер телефона    E-mail:

Вариант №2 (для индивидуальных предпринимателей)
Индивидуальный предприниматель (ФИО полностью)   ОГРНИП     ИНН
паспорт серия       №      выдан       «  »       20  , код подразделения
Адрес регистрации:    Почтовый адрес    Телефон    E-mail:

Вариант №3 (для физического лица)
(ФИО полностью)
паспорт серия       №      выдан       «  »       20  , код подразделения
Адрес регистрации:    Почтовый адрес    Телефон    E-mail:

Должник:
Вариант №1 (для юридических лиц)  (полное и краткое наименование)   ИНН
Вариант №2 (для индивидуальных предпринимателей)  Индивидуальный предприниматель (ФИО)  ИНН

Основной договор:
Вариант №1   Кредитный договор №      от «  »       20  г.
Вариант №2   Договор о предоставлении независимой гарантии №      от «  »       20 г.

Предмет залога: (описание имущества)
Место нахождения предмета залога:
Стоимость предмета залога:       руб.
Срок действия Договора залога:
--- КОНЕЦ ---
Количество вариантов документа: 1 (варианты внутри документа). Таблиц Word: 4. Приложений: 1.""",
    },
    {
        "role": "assistant",
        "content": """{
  "variables": 110,
  "variative_blocks": 15,
  "calculated_fields": 0,
  "tables": 4,
  "complexity": "Высокая",
  "description": "Индивидуальные условия договора залога движимого имущества (банк Уралсиб). Три типа залогодателя (юрлицо/ИП/физлицо), два типа должника, два типа основного договора — итого 15+ вариативных блоков по всему документу.",
  "confidence": 0.86,
  "notes": "110 переменных — оценка с учётом что поля реквизитов для каждой стороны уникальны. Высокая уверенность в вариативных блоках (маркер 'Вариант №1/2/3' явный). Рекомендуется ручная проверка переменных ±15."
}""",
    },

    # ── Пример 5: Договор с ФЛ, фиксированная оплата (Низкая, реальный документ) ──
    {
        "role": "user",
        "content": """Проанализируй документ:
--- ДОКУМЕНТ: Договор_с_ФЛ_фиксированная_оплата ---
(Инструкции для разработчика — красным курсивом. Переменные выделены цветом или обозначены [ ] и ___.)

ДОГОВОР ГРАЖДАНСКО-ПРАВОВОГО ХАРАКТЕРА №[номер договора]
г. [место заключения]                                    «___» _______ 202__ г.

[Полное наименование организации], именуемое в дальнейшем «Заказчик», в лице [должность] [ФИО подписанта], действующего на основании [документ-основание], с одной стороны, и
[ФИО физического лица], именуемый(ая) в дальнейшем «Исполнитель», паспорт серия ___ № ___, выдан ___, «___» _______ 20__ г., зарегистрированный(ая) по адресу: ___, ИНН ___, с другой стороны, заключили настоящий договор о нижеследующем:

1. ПРЕДМЕТ ДОГОВОРА
1.1. Заказчик поручает, а Исполнитель принимает на себя обязательства выполнить [описание работ/услуг] и сдать результат Заказчику.
1.2. Срок выполнения: с «___» _______ 202__ г. по «___» _______ 202__ г.

2. СТОИМОСТЬ И ПОРЯДОК РАСЧЁТОВ
2.1. Вознаграждение Исполнителя составляет ___ (_______) рублей.

Вариант А (если Исполнитель — самозанятый):
2.2. Исполнитель является плательщиком налога на профессиональный доход. Заказчик не является налоговым агентом.

Вариант Б (если Исполнитель — обычное физлицо):
2.2. Заказчик удерживает НДФЛ 13% и перечисляет страховые взносы в соответствии с НК РФ.

3. ПОРЯДОК ПРИЁМКИ
Вариант А: Результат подтверждается Актом выполненных работ, подписанным обеими сторонами.
Вариант Б: Результат подтверждается в электронном виде путём направления уведомления на e-mail.

4. ОТВЕТСТВЕННОСТЬ
Вариант А: Штраф за нарушение сроков — 0,5% от суммы за каждый день просрочки.
Вариант Б: Стороны освобождены от ответственности при форс-мажоре.

Наличие/отсутствие пункта о конфиденциальности — по выбору.
Наличие/отсутствие пункта о переходе прав на результат интеллектуальной деятельности — по выбору.
--- КОНЕЦ ---
Количество вариантов документа: 1 (варианты внутри). Таблиц Word: 1. Приложений: 2.
Стиль разметки: [ ] и ___ для переменных, Вариант А/Б для вариативных блоков, инструкции красным курсивом (не переменные), Word-комментарии справа, цветовое выделение.""",
    },
    {
        "role": "assistant",
        "content": """{
  "variables": 7,
  "variative_blocks": 5,
  "calculated_fields": 0,
  "tables": 1,
  "complexity": "Низкая",
  "description": "Договор ГПХ с физическим лицом на фиксированный результат. Минимум переменных (ФИО, паспорт, описание работ, сумма, даты). Вариативность по налоговому статусу (самозанятый/обычное ФЛ), способу приёмки, ответственности, конфиденциальности, правам на РИД.",
  "confidence": 0.90,
  "notes": "Инструкции красным курсивом — НЕ переменные, не считаются. Word-комментарии указывают на места вариативности. Специфических переменных мало — всего 7 сверх типового блока реквизитов не применимо (ФЛ не имеет ИНН организации/КПП)."
}""",
    },
]

# ── Load additional labelled examples from JSON (few_shot_labeled.json) ──────
_LABELED_JSON = os.path.join(os.path.dirname(__file__), "few_shot_labeled.json")
if os.path.exists(_LABELED_JSON):
    try:
        import json as _json
        with open(_LABELED_JSON, "r", encoding="utf-8") as _f:
            _extra = _json.load(_f)
        # Insert before existing examples so labelled real docs come first
        FEW_SHOT_EXAMPLES = _extra + FEW_SHOT_EXAMPLES
    except Exception as _e:
        print(f"[WARN] Could not load {_LABELED_JSON}: {_e}")


def build_analysis_prompt(
    main_doc: DocInfo,
    all_docs: list[DocInfo],
    diff_summary: str,
    variants_count: int,
    python_var_count: int = 0,
    python_var_raw: int = 0,
    python_block_count: int = 0,
    total_appendices: int = 0,
) -> str:
    """Assemble the user message for Claude.

    Smart extraction strategy:
      1. First 8000 chars of paragraph text (covers most documents fully)
      2. If doc is longer: extract ALL paragraphs with variative markers from the rest
         (so the LLM never misses a variative block even in long documents)
      3. Table content (first 3 tables, up to 400 chars each)
      4. All Word comments (key source of variative block info)
      5. python_var_count — pre-counted by Python; LLM must use this as final answer
    """
    # Regex for variative markers — used to extract relevant paragraphs from long docs
    _VAR_RE = re.compile(
        r'\[[^\]]{3,}\]'                           # [любой текст] в скобках
        r'|вариант\s*[№#]?\s*\d'                   # Вариант №1 / Вариант 2
        r'|вариант\s+[АБВГабвгABCDabcd][\s:.]'    # Вариант А / Вариант B
        r'|выбрать\s+вариант|выберите'
        r'|в\s+случае\s+если|при\s+наличии'
        r'|если\s+применимо|при\s+необходимости'
        r'|ненужное\s+зачеркнуть|подлежит\s+удалению'
        r'|\bили\b.{3,30}\bили\b',                 # или ... или
        re.IGNORECASE
    )

    # GPT-4o: 128K токенов. Русский текст ≈ 1 токен/3 символа.
    # Лимит GPT-4o: 128K токенов. Русский текст ≈ 1 токен / 1.5 симв.
    # Бюджет: 128K − 1K (ответ) − 6K (системный промпт) − 18K (few-shot) = ~103K токенов
    # 103K токенов × 1.5 симв/токен = ~154K симв. Берём 80K с запасом (≈53K токенов).
    # Этого хватает для 50-60 страниц плотного юртекста — достаточно для любого договора.
    FIRST_CHUNK = 80_000
    MAX_VAR_LINES = 400
    MAX_LINE_LEN  = 110

    all_para_text = "\n".join(main_doc.paragraphs)
    total_len = len(all_para_text)

    # ── Вариативные блоки: сканируем ВСЕ параграфы документа (не обрезаем!) ──────
    # Для каждого параграфа с маркером вариативности включаем его + следующие
    # N параграфов — LLM видит И инструкцию, И абзацы под её действием.
    # Это решает проблему "маркер найден, но непонятно сколько абзацев он затрагивает".
    _BLOCK_INSTR_RE = re.compile(
        r'\[(?!МАРКЕР:)[^\]]{5,200}\]'        # [инструкция в скобках]
        r'|вариант\s*[№#]?\s*\d'              # Вариант №1
        r'|вариант\s+[АБВГабвгABCDabcd][\s:.]',  # Вариант А
        re.IGNORECASE
    )
    all_paras = main_doc.paragraphs   # ВСЕ параграфы, без обрезки
    block_excerpt_parts: list[str] = []
    seen_block_paras: set[int] = set()
    CONTEXT_AFTER = 6   # сколько параграфов после маркера включать как контекст
    # Список маркерных позиций
    marker_indices: list[int] = []
    for idx, para in enumerate(all_paras):
        if _BLOCK_INSTR_RE.search(para):
            marker_indices.append(idx)
    n_markers = len(marker_indices)
    py_governed_total = 0
    for mi in marker_indices:
        # Включаем маркер + следующие CONTEXT_AFTER параграфов (без повторов)
        for j in range(mi, min(mi + 1 + CONTEXT_AFTER, len(all_paras))):
            if j not in seen_block_paras and all_paras[j].strip():
                seen_block_paras.add(j)
                block_excerpt_parts.append(all_paras[j][:MAX_LINE_LEN])

    if total_len <= FIRST_CHUNK:
        # Короткий документ — всё влезает в основной чанк
        doc_text = all_para_text
        variative_note = ""
    else:
        # Длинный документ: первые FIRST_CHUNK символов + отдельный блок для вариативности
        doc_text = all_para_text[:FIRST_CHUNK]
        trunc_chars = total_len - FIRST_CHUNK
        variative_note = (
            f"\n\n...[документ продолжается ещё {trunc_chars} символов — "
            f"переменные из обрезанной части отражены в разделе ниже]..."
        )

    # Добавляем блок вариативных маркеров из ВСЕГО документа (не только из остатка)
    if block_excerpt_parts:
        variative_note += (
            f"\n\n--- МАРКЕРЫ ВАРИАТИВНЫХ БЛОКОВ (весь документ, {n_markers} маркеров) ---\n"
            f"Каждый [маркер] показан с {CONTEXT_AFTER} следующими абзацами — "
            f"подсчитай сколько абзацев затрагивает каждое условие.\n\n"
            + "\n".join(block_excerpt_parts[:300])  # max 300 строк
        )

    # Table content (may contain variative markers too)
    table_excerpts = []
    for i, t in enumerate(main_doc.tables_content[:5]):
        t_stripped = t.strip()
        if t_stripped:
            if _VAR_RE.search(t_stripped):
                table_excerpts.append(f"[Таблица {i+1} — содержит маркеры]\n{t_stripped[:800]}")
            else:
                table_excerpts.append(f"[Таблица {i+1}]\n{t_stripped[:400]}")
    table_text = ("\n\n--- ТАБЛИЦЫ ---\n" + "\n---\n".join(table_excerpts)) if table_excerpts else ""

    # Appendix text — include paragraphs from embedded appendices so LLM can count variables there
    if main_doc.appendix_paragraphs:
        app_para_text = "\n".join(main_doc.appendix_paragraphs)
        # Respect the same 200K total budget: add up to what's left
        remaining_budget = max(0, FIRST_CHUNK - len(all_para_text))
        if remaining_budget > 2000:
            app_excerpt = app_para_text[:remaining_budget]
        else:
            # Budget tight: include only lines with variable markers
            app_var_lines = [
                ln.strip()[:MAX_LINE_LEN]
                for ln in app_para_text.split("\n")
                if ln.strip() and _VAR_RE.search(ln)
            ][:MAX_VAR_LINES]
            app_excerpt = "\n".join(app_var_lines)
        if app_excerpt.strip():
            variative_note += (
                f"\n\n--- ТЕКСТ ПРИЛОЖЕНИЙ (только для подсчёта ТАБЛИЦ и ВАРИАТИВНЫХ БЛОКОВ) ---\n"
                f"⛔ НЕ считай поля в приложениях как переменные шаблона — они заполняются вручную.\n"
                f"{app_excerpt}"
            )

    # Appendices note (file-based + in-document)
    if total_appendices > 0:
        appendix_note = f"Количество приложений: {total_appendices} (включая встроенные в документ)."
    else:
        file_appendices = [d for d in all_docs if d.is_appendix]
        if file_appendices:
            appendix_note = f"Количество приложений: {len(file_appendices)}."
        else:
            appendix_note = "Приложений нет."

    lines = [
        f"Проанализируй документ:",
        f"--- ДОКУМЕНТ: {main_doc.name} ---",
        doc_text,
        variative_note,
        table_text,
        f"--- КОНЕЦ ---",
        f"",
        f"Дополнительный контекст:",
        f"Количество вариантов документа: {variants_count}.",
        f"Содержательных таблиц Word в ОСНОВНОМ тексте договора: {main_doc.content_tables_count} "
        f"(таблицы реквизитов, подписей сторон и таблицы из приложений уже исключены). "
        f"Считай по определению ТАБЛИЦА выше.",
        f"Страниц в документе: {main_doc.pages_estimate}.",
        appendix_note,
    ]

    if main_doc.comments:
        lines.append(f"\nWord-комментарии в документе ({len(main_doc.comments)} шт.) — важный источник вариативных блоков:")
        for c in main_doc.comments[:15]:
            lines.append(f"  • {c[:150]}")

    if diff_summary:
        lines.append("\nРезультат сравнения версий документа:")
        lines.append(diff_summary[:3000])

    # Python deterministic variable count — soft upper bound for the LLM
    # НЕ директивное «ДОЛЖНО» — иначе LLM якорится к числу и не снижает.
    # Clamp ≤ py_var_count делается в коде Python уже после LLM-ответа.
    # Подсказка по переменным: показываем только если dedup > 0, используем dedup (не raw),
    # и явно говорим что это НИЖНЯЯ оценка (реальных может быть больше).
    # Число py_raw намеренно НЕ показываем — LLM якорится к нему и игнорирует few-shot.
    if python_var_count > 0:
        lines.append(
            f"\n⚠ ОРИЕНТИР ПЕРЕМЕННЫХ:\n"
            f"Python-счётчик нашёл ≥{python_var_count} уникальных позиций ___ и [МАРКЕР:] "
            f"с разным контекстом.\n"
            f"→ Считай КАЖДОЕ уникальное смысловое поле как отдельную переменную.\n"
            f"→ «Название», «представитель», «основание» для КАЖДОЙ стороны — "
            f"отдельные переменные.\n"
            f"→ ⛔ НЕ якорься к этому числу — сделай независимый анализ текста.\n"
            f"→ ⛔ Не считай подписи/реквизиты в конце — они дублируют переменные шапки.\n"
            f"→ ⛔ Не добавляй переменные которых НЕТ в тексте.\n"
            f"Это ориентир касается ТОЛЬКО поля \"variables\"."
        )

    if python_block_count > 0:
        # py_governed_total вычислен выше при сборке block_section_texts (в той же функции)
        lines.append(
            f"\n⚠ СЧЁТЧИК ВАРИАТИВНЫХ БЛОКОВ (Python, только для поля \"variative_blocks\"):\n"
            f"Скрипт нашёл {python_block_count} уникальных [скобочных условий/инструкций] в тексте.\n"
            f"Это НИЖНЯЯ ГРАНИЦА — маркеры 'Вариант А/Б', Word-комментарии скриптом не учитываются.\n"
            f"→ Поле \"variative_blocks\" ДОЛЖНО быть ≥ {python_block_count}.\n"
            f"→ Один [маркер] может управлять несколькими абзацами — считай КАЖДЫЙ управляемый абзац.\n"
            f"Это ограничение касается ТОЛЬКО поля \"variative_blocks\"."
        )

    return "\n".join(lines)


def call_llm(prompt: str, client: OpenAI) -> dict:
    """Call LLM via OpenRouter (OpenAI-compatible) and return parsed JSON."""
    # Build messages: system + few-shot pairs + current prompt.
    # Token budget: GPT-4o limit 128K tokens. Russian text ≈ 1 token per 1.5 chars.
    # Reserve 1024 for output + ~5% overhead → input budget ≈ 126K tokens ≈ 189K chars.
    CHAR_BUDGET = 180_000   # conservative chars budget for input

    sys_chars = len(FEW_SHOT_SYSTEM)
    doc_chars = len(prompt)
    overhead = sys_chars + doc_chars + 200   # 200 = role labels etc.

    # Build few-shot list: keep as many pairs as budget allows.
    # Pairs come in (user, assistant) pairs — drop from the BACK first
    # (least important: code examples come last, real labeled docs come first).
    selected_examples: list[dict] = []
    budget_left = CHAR_BUDGET - overhead
    # Walk pairs front-to-back, include if they fit
    pairs = [FEW_SHOT_EXAMPLES[i:i+2] for i in range(0, len(FEW_SHOT_EXAMPLES) - 1, 2)]
    for pair in pairs:
        pair_chars = sum(len(m["content"]) for m in pair)
        if budget_left >= pair_chars:
            selected_examples += pair
            budget_left -= pair_chars
        else:
            break   # rest won't fit; labeled docs (most important) already included

    if len(selected_examples) < len(FEW_SHOT_EXAMPLES):
        dropped = (len(FEW_SHOT_EXAMPLES) - len(selected_examples)) // 2
        print(f"    [INFO] Prompt budget: dropped {dropped} few-shot example(s) to fit context window.")

    messages = [{"role": "system", "content": FEW_SHOT_SYSTEM}]
    messages += selected_examples
    messages += [{"role": "user", "content": prompt}]

    text = ""
    try:
        response = client.chat.completions.create(
            model=MODEL,
            max_tokens=1024,   # JSON ответ занимает ~100-200 токенов; 1024 даёт запас
            temperature=0.0,   # Детерминированный вывод — важно для воспроизводимости
            messages=messages,
        )
        text = response.choices[0].message.content or ""

        # Strip markdown code fences if present
        text = re.sub(r"```(?:json)?", "", text).strip()
        result = json.loads(text)
        return result

    except json.JSONDecodeError as e:
        print(f"  [WARN] JSON parse error: {e}. Raw: {text[:200]}")
        return {"error": str(e), "raw": text}
    except Exception as e:
        print(f"  [ERROR] LLM API error: {e}")
        return {"error": str(e)}


# ─────────────────────────────────────────────
# STEP 4 — Group documents and run analysis
# ─────────────────────────────────────────────

def group_documents(docs: list[DocInfo]) -> list[dict]:
    """
    Group documents into logical sets for analysis.
    Strategy:
    - All non-appendix docs form one "main" group (they are variants of the same template).
    - Each appendix is analyzed separately.
    Returns list of {group_name, main_docs, appendix_docs, pairs}.
    """
    main_docs = [d for d in docs if not d.is_appendix]
    appendix_docs = [d for d in docs if d.is_appendix]

    groups = []

    # Main documents group
    if main_docs:
        # Build comparison pairs (cross-product of short names)
        pairs = _build_pairs(main_docs)
        groups.append({
            "group_name": "Основной договор",
            "main_docs": main_docs,
            "appendix_docs": appendix_docs,
            "pairs": pairs,
            "variants_count": len(main_docs),
        })

    # Each appendix separately
    for app in appendix_docs:
        groups.append({
            "group_name": f"Приложение: {app.name}",
            "main_docs": [app],
            "appendix_docs": [],
            "pairs": [],
            "variants_count": 1,
        })

    return groups


def _build_pairs(docs: list[DocInfo]) -> list[tuple[DocInfo, DocInfo]]:
    """Build meaningful comparison pairs from a list of documents."""
    if len(docs) < 2:
        return []

    # Try to pair by similar name (EXW↔DAP, one↔multiple)
    def similarity_key(name: str) -> str:
        name = name.lower()
        name = re.sub(r"(exw|dap)", "BASIS", name)
        name = re.sub(r"(множественность|множество|несколько)", "MULTI", name)
        name = re.sub(r"\s+", " ", name).strip()
        return name

    # Group by base name
    from itertools import combinations
    pairs = []
    for a, b in combinations(docs, 2):
        # Pair if names are similar enough (share >50% of tokens)
        tokens_a = set(re.split(r"[\s_\-]+", a.name.lower()))
        tokens_b = set(re.split(r"[\s_\-]+", b.name.lower()))
        overlap = len(tokens_a & tokens_b) / max(len(tokens_a | tokens_b), 1)
        if overlap > 0.4:
            pairs.append((a, b))

    # If no pairs found, just pair first with second
    if not pairs and len(docs) >= 2:
        pairs.append((docs[0], docs[1]))

    return pairs[:4]  # Max 4 pairs to keep prompt size manageable


# ─────────────────────────────────────────────
# STEP 5 — Write to Excel
# ─────────────────────────────────────────────

# Column mapping for "Анализ сложности" sheet
EXCEL_COLUMNS = {
    "Порядковый номер": 1,
    "Номер основного документа": 2,
    "Название документа": 3,
    "Количество вариантов": 4,
    "Кол-во страниц": 5,
    "Кол-во приложений": 6,
    "Кол-во переменных": 7,
    "Кол-во вариативных пунктов": 8,
    "Кол-во расчитываемых полей": 9,
    "Кол-во таблиц": 10,
    "Аналитика": 11,
    "Описание": 12,
    "Сложность": 13,
    # Коэф (14) — formula
    # Анализ исх. документа (15) — formula
    # Встречи (16) — formula
    # Тестирование (17) — formula
}

COMPLEXITY_COEF = {"Высокая": 1.5, "Средняя": 1.0, "Низкая": 0.8}


def write_to_excel(results: list[AnalysisResult], excel_path: str):
    """Append rows to the 'Анализ сложности' sheet in the Excel file."""
    try:
        wb = openpyxl.load_workbook(excel_path)
    except FileNotFoundError:
        print(f"  [WARN] Excel file not found: {excel_path}. Creating new file.")
        wb = openpyxl.Workbook()
        wb.active.title = "Анализ сложности"

    # Find the sheet
    sheet_name = "Анализ сложности"
    if sheet_name not in wb.sheetnames:
        print(f"  [WARN] Sheet '{sheet_name}' not found. Available: {wb.sheetnames}")
        ws = wb.create_sheet(sheet_name)
        # Write headers
        for name, col in EXCEL_COLUMNS.items():
            ws.cell(row=1, column=col, value=name)
        start_row = 2
    else:
        ws = wb[sheet_name]
        # Find first empty row after header
        start_row = ws.max_row + 1
        if start_row <= 2:
            start_row = 2

    for i, res in enumerate(results, start=1):
        row = start_row + i - 1
        coef = COMPLEXITY_COEF.get(res.complexity, 1.0)
        ws.cell(row=row, column=1, value=row - 1)           # Порядковый номер
        ws.cell(row=row, column=2, value=1)                  # Номер осн. документа
        ws.cell(row=row, column=3, value=res.doc_name)
        ws.cell(row=row, column=4, value=res.variants_count)
        ws.cell(row=row, column=5, value=res.pages)
        ws.cell(row=row, column=6, value=res.appendices)
        ws.cell(row=row, column=7, value=res.variables)
        ws.cell(row=row, column=8, value=res.variative_blocks)
        ws.cell(row=row, column=9, value=res.calculated_fields)
        ws.cell(row=row, column=10, value=res.tables)
        ws.cell(row=row, column=11, value="")                # Аналитика (ч) — manual
        ws.cell(row=row, column=12, value=res.description)
        ws.cell(row=row, column=13, value=res.complexity)
        # Note: columns 14-17 (Коэф, Анализ, Встречи, Тестирование) are formula-driven
        # in the original Excel — we leave them to be calculated automatically.

    wb.save(excel_path)
    print(f"  Excel updated: {excel_path} (rows {start_row}–{start_row + len(results) - 1})")


# ─────────────────────────────────────────────
# STEP 6 — Main orchestration
# ─────────────────────────────────────────────

def analyse_group(
    g: dict,
    client,
    all_docs: list | None = None,
    dry_run: bool = False,
    debug: bool = False,
) -> "AnalysisResult | None":
    """Analyse one document group and return AnalysisResult (or None on error)."""
    group_name = g["group_name"]
    main_docs: list[DocInfo] = g["main_docs"]
    appendix_docs: list[DocInfo] = g["appendix_docs"]
    pairs = g["pairs"]
    variants_count = g["variants_count"]
    docs = all_docs or main_docs

    diff_summary = build_diff_summary(pairs) if pairs else ""
    rep_doc = max(main_docs, key=lambda d: len(d.paragraphs))
    total_file_appendices = len(appendix_docs) + rep_doc.in_doc_appendices

    print(f"  Analysing group: '{group_name}' (representative: {rep_doc.name})")
    print(f"    Tables: {rep_doc.tables_count} всего, "
          f"{rep_doc.content_tables_count} содержательных в основном теле "
          f"(отфильтровано {rep_doc.tables_count - rep_doc.content_tables_count} "
          f"служебных/в-приложениях), встроенных приложений: {rep_doc.in_doc_appendices}")

    if debug:
        content_set = set(rep_doc.content_tables_content)
        print(f"\n  [DEBUG] Table classification ({rep_doc.name}):")
        for i, tbl in enumerate(rep_doc.tables_content):
            label = "СОДЕРЖАТЕЛЬНАЯ ✓" if tbl in content_set else "служебная (filtered)"
            rows = tbl.split('\n')
            preview = " // ".join(r[:60] for r in rows[:3])
            print(f"    Таблица {i+1} [{label}]: {preview[:200]}")
        print()

    py_var_count, py_var_raw, var_debug_lines = count_variables_deterministic(
        rep_doc.paragraphs,
        rep_doc.content_tables_content
    )
    print(f"    Python var count: {py_var_count} deduped / {py_var_raw} raw")
    if debug:
        print("\n".join(var_debug_lines[:50]))

    py_block_count, block_debug_lines = count_variative_python(rep_doc.paragraphs)
    print(f"    Python block count: {py_block_count} (нижняя граница)")
    if debug:
        print("\n".join(block_debug_lines[:40]))

    if dry_run:
        llm_result = {
            "variables": py_var_count if py_var_count > 0 else 15,
            "variative_blocks": 4,
            "calculated_fields": 2,
            "tables": rep_doc.content_tables_count,
            "complexity": "Средняя",
            "description": "[DRY RUN] Dummy analysis result",
            "confidence": 0.0,
            "notes": "Dry run — no API call made",
        }
    else:
        prompt = build_analysis_prompt(
            rep_doc, docs, diff_summary, variants_count,
            python_var_count=py_var_count,
            python_var_raw=py_var_raw,
            python_block_count=py_block_count,
            total_appendices=total_file_appendices,
        )
        print(f"    Calling Claude... (prompt ≈ {len(prompt)} chars)")
        llm_result = call_llm(prompt, client)

    if "error" in llm_result:
        print(f"    [WARN] LLM error: {llm_result['error']}")
        confidence = 0.0
    else:
        confidence = llm_result.get("confidence", 0.5)
        llm_vars_raw = llm_result.get("variables", 0)

        if py_var_count > 0 and py_var_count <= 30 and llm_vars_raw > py_var_count + 2:
            final_vars = py_var_count + 2
            print(f"    [CLAMP] vars: {llm_vars_raw} → {final_vars} (py_dedup={py_var_count}, soft-cap=py+2)")
        else:
            final_vars = llm_vars_raw
        llm_result["variables"] = final_vars

        py_tables = rep_doc.content_tables_count
        llm_tables = llm_result.get("tables", rep_doc.content_tables_count)
        if py_tables > 0:
            clamped_tables = min(llm_tables, py_tables)
            if clamped_tables != llm_tables:
                print(f"    [CLAMP tables] LLM={llm_tables} → {clamped_tables} (Python max={py_tables})")
                llm_result["tables"] = clamped_tables
            elif llm_tables > py_tables:
                print(f"    [CLAMP tables] LLM={llm_tables} > Python={py_tables} → {py_tables}")
                llm_result["tables"] = py_tables
        elif py_tables == 0 and llm_tables > 0:
            print(f"    [CLAMP tables] Python=0 → forcing 0 (LLM={llm_tables})")
            llm_result["tables"] = 0

        if py_block_count > 0:
            llm_blocks = llm_result.get("variative_blocks", 0)
            if llm_blocks < py_block_count:
                print(f"    [CLAMP blocks] LLM={llm_blocks} → {py_block_count} (Python lower bound)")
                llm_result["variative_blocks"] = py_block_count

        print(f"    Done. confidence={confidence:.2f}, "
              f"variables={llm_result.get('variables')} "
              f"(llm={llm_vars_raw}, py_dedup={py_var_count}, py_raw={py_var_raw}), "
              f"variative_blocks={llm_result.get('variative_blocks')}, "
              f"complexity={llm_result.get('complexity')}")

    final_vars = llm_result.get("variables", 0)
    final_blocks = llm_result.get("variative_blocks", 0)
    final_calc = llm_result.get("calculated_fields", 0)
    if final_vars > 60 or final_blocks > 35 or final_calc > 0:
        final_complexity = "Высокая"
    elif final_vars >= 30 or final_blocks >= 10:
        final_complexity = "Средняя"
    else:
        final_complexity = "Низкая"
    llm_complexity = llm_result.get("complexity", "Средняя")
    if final_complexity != llm_complexity:
        print(f"    [COMPLEXITY] LLM={llm_complexity} → recalc={final_complexity} "
              f"(vars={final_vars}, blocks={final_blocks}, calc={final_calc})")

    return AnalysisResult(
        doc_name=group_name,
        variants_count=variants_count,
        pages=rep_doc.pages_estimate,
        appendices=total_file_appendices,
        variables=final_vars,
        variative_blocks=final_blocks,
        calculated_fields=final_calc,
        tables=llm_result.get("tables", rep_doc.content_tables_count),
        complexity=final_complexity,
        doc_type=llm_result.get("doc_type", ""),
        description=llm_result.get("description", ""),
        confidence=confidence,
        raw_llm=llm_result,
        found_variables=llm_result.get("found_variables", []),
        found_blocks=llm_result.get("found_blocks", []),
        found_tables=llm_result.get("found_tables", []),
        found_calculated=llm_result.get("found_calculated", []),
    )


def run_analysis(docs_dir: str, excel_path: str, dry_run: bool = False, debug: bool = False):
    """Full pipeline: read → compare → analyse → write."""
    print(f"\n{'='*60}")
    print(f"TurboContract Document Analysis Agent")
    print(f"{'='*60}")
    print(f"Docs dir : {docs_dir}")
    print(f"Excel    : {excel_path}")
    print(f"Model    : {MODEL}")
    print(f"Dry run  : {dry_run}")
    print(f"Debug    : {debug}")
    print()

    # 1. Read all .docx files
    print("[1/4] Reading documents...")
    docx_files = list(Path(docs_dir).glob("*.docx"))
    if not docx_files:
        print(f"  No .docx files found in {docs_dir}")
        return

    docs: list[DocInfo] = []
    for f in sorted(docx_files):
        print(f"  Reading: {f.name}")
        try:
            info = extract_docx(str(f))
            docs.append(info)
            print(f"    → {len(info.paragraphs)} paragraphs, {info.tables_count} tables, "
                  f"{len(info.comments)} comments, ~{info.pages_estimate} pages "
                  f"(words={info.word_count}), appendix={info.is_appendix}")
            if debug:
                # Print first 2000 chars of extracted text so user can verify [МАРКЕР:...] tags
                preview_text = "\n".join(info.paragraphs)[:2000]
                print(f"\n  [DEBUG] Extracted text preview (first 2000 chars) — {f.name}:")
                print("  " + preview_text.replace("\n", "\n  "))
                print("  [DEBUG] ...\n")
        except Exception as e:
            print(f"  [ERROR] Failed to read {f.name}: {e}")

    print(f"  Total: {len(docs)} documents loaded.\n")

    # 2. Group and build pairs
    print("[2/4] Grouping documents...")
    groups = group_documents(docs)
    for g in groups:
        print(f"  Group: '{g['group_name']}' | variants={g['variants_count']} | "
              f"pairs={len(g['pairs'])}")
    print()

    # 3. Build diff summaries and call LLM via OpenRouter
    print("[3/4] Analysing with LLM (OpenRouter)...")
    if not OPENROUTER_API_KEY and not dry_run:
        print("  [ERROR] OPENROUTER_API_KEY not set. Use --dry-run to skip LLM calls.")
        print("  Задай ключ: set OPENROUTER_API_KEY=sk-or-v1-...")
        return

    client = OpenAI(
        api_key=OPENROUTER_API_KEY,
        base_url=OPENROUTER_BASE_URL,
    ) if not dry_run else None

    all_results: list[AnalysisResult] = []

    for g in groups:
        result = analyse_group(g, client, all_docs=docs, dry_run=dry_run, debug=debug)
        if result:
            all_results.append(result)
            if result.confidence < 0.7:
                print(f"    ⚠  LOW CONFIDENCE ({result.confidence:.2f}) — human review recommended")

    print()

    # 4. Write to Excel
    print("[4/4] Writing to Excel...")
    if not dry_run:
        try:
            write_to_excel(all_results, excel_path)
        except Exception as e:
            print(f"  [ERROR] Excel write failed: {e}")
    else:
        print("  (Dry run — skipping Excel write)")

    # 5. Print summary
    print()
    print(f"{'='*60}")
    print("SUMMARY")
    print(f"{'='*60}")
    for res in all_results:
        flag = "⚠ " if res.confidence < 0.7 else "✓ "
        print(f"{flag}{res.doc_name}")
        if res.doc_type:
            print(f"   Тип: {res.doc_type}")
        print(f"   Варианты={res.variants_count}  Страниц≈{res.pages}  "
              f"Приложений={res.appendices}")
        print(f"   Переменных={res.variables}  Вариативных={res.variative_blocks}  "
              f"Расчётных={res.calculated_fields}  Таблиц={res.tables}")
        print(f"   Сложность={res.complexity}  Уверенность={res.confidence:.0%}")
        if res.description:
            print(f"   Описание: {res.description[:100]}")
        # ── Детализация найденного ──────────────────────────────────────────
        if res.found_variables:
            print(f"\n   📋 ПЕРЕМЕННЫЕ ({len(res.found_variables)}):")
            for i, v in enumerate(res.found_variables, 1):
                print(f"      {i:2}. {v}")
        if res.found_blocks:
            print(f"\n   🔀 ВАРИАТИВНЫЕ БЛОКИ ({len(res.found_blocks)}):")
            for i, b in enumerate(res.found_blocks, 1):
                print(f"      {i:2}. {b}")
        if res.found_tables:
            print(f"\n   📊 ТАБЛИЦЫ ({len(res.found_tables)}):")
            for i, t in enumerate(res.found_tables, 1):
                print(f"      {i:2}. {t}")
        if res.found_calculated:
            print(f"\n   🧮 РАСЧЁТНЫЕ ПОЛЯ ({len(res.found_calculated)}):")
            for i, c in enumerate(res.found_calculated, 1):
                print(f"      {i:2}. {c}")
        print()

    # Dump full JSON for review
    json_out = Path(docs_dir) / "analysis_result.json"
    with open(json_out, "w", encoding="utf-8") as f:
        json.dump(
            [asdict(r) for r in all_results],
            f,
            ensure_ascii=False,
            indent=2,
        )
    print(f"Full JSON saved: {json_out}")

    # 6. Save feedback template for human correction (future RAG corpus)
    if not dry_run:
        _save_feedback_template(all_results, docs, docs_dir)


def _save_feedback_template(
    results: list[AnalysisResult],
    docs: list[DocInfo],
    docs_dir: str,
):
    """
    Save / update feedback_library.json — a file where the user can fill in
    the correct values after manual review. Over time this becomes the RAG corpus.

    Structure per entry:
      doc_name    — group name (e.g. "Основной договор")
      doc_excerpt — first 600 chars of representative document text (for identification)
      auto        — values produced by the agent this run
      correct     — null fields for the user to fill in manually
                    (once filled, these entries can be used as few-shot examples / RAG)
    """
    library_path = Path(docs_dir) / "feedback_library.json"

    # Load existing library (append-mode — don't overwrite previous corrections)
    existing: list[dict] = []
    if library_path.exists():
        try:
            with open(library_path, "r", encoding="utf-8") as f:
                existing = json.load(f)
        except Exception:
            existing = []

    # Build a lookup of already-corrected entries by doc_name (skip re-adding)
    existing_names = {e["doc_name"] for e in existing if e.get("correct", {}).get("variables") is not None}

    # Build doc_name → representative DocInfo mapping
    doc_map: dict[str, DocInfo] = {}
    for d in docs:
        if not d.is_appendix:
            doc_map.setdefault("Основной договор", d)
            # Keep the largest doc as representative
            if len(d.paragraphs) > len(doc_map.get("Основной договор", d).paragraphs):
                doc_map["Основной договор"] = d
    for d in docs:
        if d.is_appendix:
            doc_map[f"Приложение: {d.name}"] = d

    new_entries = 0
    for res in results:
        if res.doc_name in existing_names:
            # Already has a human correction — don't overwrite
            continue

        rep = doc_map.get(res.doc_name)
        excerpt = ""
        if rep:
            excerpt = "\n".join(rep.paragraphs)[:600].strip()

        entry = {
            "doc_name": res.doc_name,
            "doc_excerpt": excerpt,
            "auto": {
                "variables": res.variables,
                "variative_blocks": res.variative_blocks,
                "calculated_fields": res.calculated_fields,
                "tables": res.tables,
                "pages": res.pages,
                "complexity": res.complexity,
            },
            "correct": {
                "variables": None,       # ← заполни правильное число
                "variative_blocks": None,
                "calculated_fields": None,
                "tables": None,
                "pages": None,
                "complexity": None,
            },
            "notes": "",  # ← необязательный комментарий
        }

        # Replace existing uncorrected entry for this doc_name (update auto values)
        replaced = False
        for i, e in enumerate(existing):
            if e["doc_name"] == res.doc_name:
                existing[i] = entry
                replaced = True
                break
        if not replaced:
            existing.append(entry)
        new_entries += 1

    with open(library_path, "w", encoding="utf-8") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)

    uncorrected = sum(
        1 for e in existing
        if e.get("correct", {}).get("variables") is None
    )
    corrected = len(existing) - uncorrected
    print(f"Feedback library: {library_path}")
    print(f"  Всего записей: {len(existing)}  "
          f"(исправлено человеком: {corrected}, ожидают проверки: {uncorrected})")
    if uncorrected > 0:
        print(f"  → Открой feedback_library.json и заполни поля 'correct' → null замени на числа.")


# ─────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="TurboContract Document Analysis Agent"
    )
    parser.add_argument(
        "--docs",
        default=DEFAULT_DOCS_DIR,
        help="Path to folder with .docx files",
    )
    parser.add_argument(
        "--excel",
        default=DEFAULT_EXCEL,
        help="Path to Excel complexity table",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Run without calling Claude API or writing to Excel",
    )
    parser.add_argument(
        "--debug",
        action="store_true",
        help=(
            "Print extracted document text (first 2000 chars) and variable "
            "debug info so you can verify [МАРКЕР:...] tags are detected correctly"
        ),
    )
    args = parser.parse_args()

    run_analysis(
        docs_dir=args.docs,
        excel_path=args.excel,
        dry_run=args.dry_run,
        debug=args.debug,
    )


if __name__ == "__main__":
    main()
